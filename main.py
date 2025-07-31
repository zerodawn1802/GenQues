import json
import re
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image
from io import BytesIO
import requests
import validators
from google import genai
from google.genai import types
from text2Image import generate_image_from_text

def parse_markdown_table(lines):
    headers = []
    rows = []
    for i, line in enumerate(lines):
        cells = [cell.strip() for cell in re.split(r'(?<!\\)\|', line)[1:-1]]
        if i == 0:
            headers = cells
        elif i == 1:
            if all(re.match(r'^[:]?[-]+[:]?\s*$', cell) for cell in cells):
                is_table = True
            else:
                print(f"Invalid separator line: {cells}")
                return None, None
        else:
            rows.append(cells)
    return headers, rows

def add_table_to_doc(doc, headers, rows):
    if not headers or not rows:
        return
    
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    
    for i, column in enumerate(table.columns):
        column.width = Inches(2.5 if i == 1 else 1.5)
    
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
    
    # Add data rows
    for row_idx, row in enumerate(rows, 1):
        for col_idx, cell_data in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            if '<br>' in cell_data:
                lines = cell_data.split('<br>')
                for line in lines:
                    p = cell.add_paragraph(line.strip())
                    for run in p.runs:
                        run.font.size = Pt(11)
            else:
                cell.text = cell_data
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)

def create_docx_from_json(json_data, output_file):
    doc = Document()
    
    data = json.loads(json_data)
    
    for idx, item in enumerate(data, 1):
        p = doc.add_paragraph(f"Câu {idx}.")
        for run in p.runs:
            run.font.size = Pt(11)
        
        mo_ta_anh = item.get("Mô tả ảnh", "Không có")
        if mo_ta_anh != "Không có":
            try:
                if validators.url(mo_ta_anh):
                    response = requests.get(mo_ta_anh)
                    response.raise_for_status()
                    image = Image.open(BytesIO(response.content))
                else:
                    image = generate_image_from_text(mo_ta_anh)
                
                img_byte_arr = BytesIO()
                image.save(img_byte_arr, format='PNG')
                img_byte_arr.seek(0)
                doc.add_picture(img_byte_arr, width=Inches(5.0))
            except Exception as e:
                doc.add_paragraph(f"[Không thể tạo hoặc tải ảnh: {str(e)}]")
        
        noi_dung_cau_hoi = item.get("Nội dung câu hỏi", "")
        noi_dung_cau_hoi = noi_dung_cau_hoi.replace('\r\n', '\n')
        lines = noi_dung_cau_hoi.split('\n')
        table_lines = []
        table_p = False
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('|') and line.endswith('|'):
                table_lines.append(line)
                table_p = True
            else:
                if table_p:
                    headers, rows = parse_markdown_table(table_lines)
                    add_table_to_doc(doc, headers, rows)
                    table_lines = []
                    table_p = False
                p = doc.add_paragraph(line)
        doc.add_paragraph()
    
    doc.save(output_file)

# JSON data provided by the user
with open(r'C:\Users\Admin\Desktop\Maru\GenQues\result.json', 'r', encoding='utf-8') as f:
    json_data = f.read()

# Call the function to create the DOCX file
create_docx_from_json(json_data, "questions.docx")